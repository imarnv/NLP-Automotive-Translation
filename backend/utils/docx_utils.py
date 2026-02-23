import re
import os
from copy import deepcopy
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsmap
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

# 1. Global Namespace Registration
if 'v' not in nsmap:
    nsmap['v'] = "urn:schemas-microsoft-com:vml"
if 'w' not in nsmap:
    nsmap['w'] = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
if 'wp' not in nsmap:
    nsmap['wp'] = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _is_in_table(p_elem):
    parent = p_elem.getparent()
    while parent is not None:
        if parent.tag == qn('w:tc'):
            return True
        parent = parent.getparent()
    return False


def translate_docx(input_path, output_path, translation_helper, target_lang, progress_callback=None):
    if progress_callback: progress_callback("Loading document...", 10)
    doc = Document(input_path)

    # 2. Extract Segments
    segments = []
    processed_elements = set()

    for para in doc.paragraphs:
        if para.text.strip() and para._element not in processed_elements:
            segments.append({"text": para.text, "obj": para, "type": "paragraph"})
            processed_elements.add(para._element)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() and para._element not in processed_elements:
                        segments.append({"text": para.text, "obj": para, "type": "paragraph"})
                        processed_elements.add(para._element)

    for tb in doc.element.findall('.//' + qn('v:textbox')):
        for p in tb.findall('.//' + qn('w:p')):
            if p not in processed_elements:
                para = Paragraph(p, doc)
                if para.text.strip():
                    segments.append({"text": para.text, "obj": para, "type": "paragraph"})
                    processed_elements.add(p)

    if not segments:
        doc.save(output_path)
        return

    # 3. Batch Translate
    texts_to_translate = [s["text"] for s in segments]
    if progress_callback: progress_callback(f"Translating {len(segments)} segments...", 40)
    translated_texts = translation_helper(texts_to_translate, target_lang, progress_callback)

    # 4. Apply Translations & Styling
    if progress_callback: progress_callback("Reconstructing document layout...", 80)
    for i, seg in enumerate(segments):
        if i >= len(translated_texts): break
        
        para = seg["obj"]
        original_text = seg["text"]
        translated_text = translated_texts[i]
        if not translated_text: continue

        in_table = _is_in_table(para._element)
        is_in_shape = False
        if para._element.xpath('.//ancestor::w:txbxContent') or para._element.xpath('.//ancestor::v:textbox'):
            is_in_shape = True

        # 5. Capture Original Style
        target_run = None
        if para.runs:
            for r in para.runs:
                if r.text.strip():
                    target_run = r
                    break
            if not target_run: target_run = para.runs[0]
        else:
            target_run = para.add_run()

        is_bold = target_run.bold
        is_italic = target_run.italic
        underline = target_run.underline
        font_name = target_run.font.name
        font_size = target_run.font.size
        font_size_pts = font_size.pt if font_size else 11.0
        
        color_rgb = None
        theme_color = None
        if target_run.font.color:
            color_rgb = target_run.font.color.rgb
            theme_color = target_run.font.color.theme_color

        # 6. Balanced Spacing (SMART — preserve heading hierarchy)
        growth_factor = len(translated_text) / len(original_text) if original_text else 1.0

        # 7. Header Exemption Check (done BEFORE spacing so we can use it)
        is_heading = False
        if font_size_pts >= 14 or (color_rgb and str(color_rgb) != '000000'):
            is_heading = True
        style_name = (para.style.name or '').lower()
        if 'heading' in style_name or 'title' in style_name:
            is_heading = True

        if is_heading:
            # Headings: keep their original space_before (breathing room above)
            # but allow keep_with_next so they stay glued to their content.
            # Do NOT touch space_before — it's what makes headings look professional.
            para.paragraph_format.widow_control = False
            # Reduce space_after only if extreme
            if para.paragraph_format.space_after and para.paragraph_format.space_after.pt > 18:
                para.paragraph_format.space_after = Pt(12)
        else:
            # Body text: allow free flow across pages
            para.paragraph_format.widow_control = False
            para.paragraph_format.keep_with_next = False
            para.paragraph_format.keep_together = False
            # Only reduce space_after if it's huge
            if para.paragraph_format.space_after and para.paragraph_format.space_after.pt > 12:
                para.paragraph_format.space_after = Pt(6)
            # Tighten line spacing if text grew a lot
            if growth_factor > 1.5:
                para.paragraph_format.line_spacing = 1.0

        # 8. Font Scaling (Skip for headings and tables to keep them stable)
        scaled_font_size = None
        if not is_heading and not in_table and (growth_factor > 1.1 or is_in_shape):
            reduction = min(0.25, max(0.1, (growth_factor - 1.0) * 0.5))
            if is_in_shape: reduction = max(reduction, 0.2)
            new_size = font_size_pts * (1.0 - reduction)
            scaled_font_size = Pt(max(7, new_size))

        # 9. Replacement
        translated_text = translated_text.rstrip().replace('\t', ' ')
        parts = re.split(r'(@@.*?@@)', translated_text)
        text_put = False
        for run in para.runs:
            t_elements = run._element.findall(qn('w:t'))
            if not t_elements: continue
            if not text_put:
                for t_elem in t_elements: run._element.remove(t_elem)
                for idx, part in enumerate(parts):
                    if not part: continue
                    is_h = part.startswith('@@') and part.endswith('@@')
                    disp = part[2:-2] if is_h else part
                    if idx == 0:
                        run.text = disp
                        if font_name: run.font.name = font_name
                        if scaled_font_size: run.font.size = scaled_font_size
                        if is_h: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        new_r = para.add_run(disp)
                        new_r.bold = is_bold
                        new_r.italic = is_italic
                        new_r.underline = underline
                        if font_name: new_r.font.name = font_name
                        if scaled_font_size: new_r.font.size = scaled_font_size
                        if is_h: new_r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        if color_rgb: new_r.font.color.rgb = color_rgb
                        elif theme_color: new_r.font.color.theme_color = theme_color
                text_put = True
            else:
                for t_elem in t_elements: t_elem.text = ""

    # =====================================================================
    #  POST-TRANSLATION CLEANUP (REFINED)
    # =====================================================================
    body = doc.element.body

    # 10. Table row heights (Fix cropping)
    for table in doc.tables:
        for row in table.rows:
            row.allow_break_across_pages = True
            if row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY:
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # 11. Remove Page-Spill Blockers (Only what's absolutely necessary)
    
    # 11a. Remove paragraph pageBreakBefore
    for pbb in list(body.iter(qn('w:pageBreakBefore'))):
        pbb.getparent().remove(pbb)

    # 11b. Remove inline page breaks
    for br in list(body.iter(qn('w:br'))):
        if br.get(qn('w:type')) == 'page':
            parent = br.getparent()
            if parent is not None: parent.remove(br)

    # 11c. Remove stale rendered breaks (Crucial for spill-over flow)
    for lrpb in list(body.iter(qn('w:lastRenderedPageBreak'))):
        lrpb.getparent().remove(lrpb)

    # 11d. Convert ALL intermediate section breaks to CONTINUOUS
    #      CRITICAL: sectPr with NO <w:type> child defaults to 'nextPage'!
    #      We must explicitly add/set 'continuous' on every intermediate
    #      sectPr, not just ones that already have a type element.
    all_sect_prs = list(body.iter(qn('w:sectPr')))
    # The very last sectPr is the document body's final section — leave it
    # alone so page size, margins, headers/footers are preserved.
    intermediate_sects = all_sect_prs[:-1] if len(all_sect_prs) > 1 else []
    for sectPr in intermediate_sects:
        sect_type = sectPr.find(qn('w:type'))
        if sect_type is None:
            # No type element = implicit nextPage. Add explicit continuous.
            sect_type = etree.SubElement(sectPr, qn('w:type'))
        sect_type.set(qn('w:val'), 'continuous')

    # 12. Fix anchored images — minimal safe normalization
    #     - allowOverlap so images coexist in tight multi-column layouts
    #     - distT/distB breathing room so translated text doesn't crowd images
    #     - positionV: page/margin → paragraph with offset=0 so images follow
    #       their anchor paragraph after text reflow
    #     - positionH: intentionally untouched (original horizontal coords are correct)
    #     - layoutInCell: intentionally untouched
    EMU_1MM = 91440   # 1 mm in EMUs

    for anchor in body.iter(qn('wp:anchor')):
        anchor.set('allowOverlap', '1')
        anchor.set('simplePos', '0')

        # Minimum top/bottom breathing room
        try:
            if int(anchor.get('distT', '0')) < EMU_1MM:
                anchor.set('distT', str(EMU_1MM))
            if int(anchor.get('distB', '0')) < EMU_1MM:
                anchor.set('distB', str(EMU_1MM))
        except (ValueError, TypeError):
            anchor.set('distT', str(EMU_1MM))
            anchor.set('distB', str(EMU_1MM))

        # Vertical: convert page/margin-relative to paragraph-relative
        # so images follow their anchor paragraph across pages after reflow.
        # Reset offset to 0 — page-coordinate EMU values are not reusable
        # as paragraph-relative offsets.
        pos_v = anchor.find(qn('wp:positionV'))
        if pos_v is not None:
            if pos_v.get('relativeFrom', '') in ('page', 'margin'):
                pos_v.set('relativeFrom', 'paragraph')
                off = pos_v.find(qn('wp:posOffset'))
                if off is not None:
                    off.text = '0'

    # 13. Remove false vertical lines
    #     The dashed vertical lines come from TWO sources:
    #     (a) Paragraph left/right borders (w:pBdr) — sidebar markers
    #     (b) Table cell borders (w:tcBorders) — layout table dividers
    #         that become visible as overlapping lines after reflow.

    # 13a. Paragraph borders (ALL sides — left, right, top, bottom, between, bar)
    #      These produce stray horizontal and vertical lines after text reflow.
    for pBdr in list(body.iter(qn('w:pBdr'))):
        for side in (qn('w:left'), qn('w:right'), qn('w:bar'),
                     qn('w:top'), qn('w:bottom'), qn('w:between')):
            border_elem = pBdr.find(side)
            if border_elem is not None:
                pBdr.remove(border_elem)

    # 13b. Table cell borders — hide ALL visible borders that produce stray lines.
    #      After text reflow these overlap with content and images.
    #      Keep 'none' and 'nil' as-is; everything else → nil.
    KEEP_VALS = {'none', 'nil'}
    for tcBdr in list(body.iter(qn('w:tcBorders'))):
        for side in (qn('w:left'), qn('w:right'), qn('w:insideV'),
                     qn('w:top'), qn('w:bottom'), qn('w:insideH')):
            border_elem = tcBdr.find(side)
            if border_elem is not None:
                val = border_elem.get(qn('w:val'), 'none')
                if val not in KEEP_VALS:
                    border_elem.set(qn('w:val'), 'nil')

    # 13c. Table-level borders — hide inside vertical and horizontal borders
    for tblBdr in list(body.iter(qn('w:tblBorders'))):
        for side in (qn('w:insideV'), qn('w:insideH')):
            border_elem = tblBdr.find(side)
            if border_elem is not None:
                val = border_elem.get(qn('w:val'), 'none')
                if val not in KEEP_VALS:
                    border_elem.set(qn('w:val'), 'nil')

    # 13d. VML line elements (v:line) — the PRIMARY source of stray lines.
    #      These are absolutely-positioned vector lines from the original
    #      document layout.  After text reflow they stay at their original
    #      coordinates, cutting through content at wrong positions.
    #      Diagnostic found 404 v:line elements in a single translated DOCX.
    _VML_LINE = '{urn:schemas-microsoft-com:vml}line'
    _VML_STROKE = '{urn:schemas-microsoft-com:vml}stroke'
    for vline in list(body.iter(_VML_LINE)):
        # Make the line invisible without removing the element (preserves
        # any child elements or references Word expects)
        vline.set('stroked', 'false')
        vline.set('strokecolor', 'white')
        vline.set('strokeweight', '0')
        # Also neutralize any v:stroke child element
        stroke_child = vline.find(_VML_STROKE)
        if stroke_child is not None:
            vline.remove(stroke_child)

    # 13e. VML rect elements (v:rect) — bordered rectangles that
    #      appear as stray outlines after reflow.
    #      NOTE: v:shape elements are intentionally NOT touched — they form
    #      actual diagrams (flow charts, connector arrows, labelled components).
    _VML_RECT = '{urn:schemas-microsoft-com:vml}rect'
    for vrect in list(body.iter(_VML_RECT)):
        vrect.set('stroked', 'false')
        stroke_child = vrect.find(_VML_STROKE)
        if stroke_child is not None:
            vrect.remove(stroke_child)

    # 14. Fix column separators
    for sectPr in intermediate_sects:
        cols = sectPr.find(qn('w:cols'))
        if cols is not None:
            if cols.get(qn('w:sep')):
                del cols.attrib[qn('w:sep')]

    if progress_callback: progress_callback("Saving final document...", 95)
    doc.save(output_path)
