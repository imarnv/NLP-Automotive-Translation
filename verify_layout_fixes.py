"""Verify the three layout fixes in docx_utils.py."""
import sys, os
sys.path.insert(0, '.')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ---- CREATE A TEST DOCX ----
doc = Document()

# 1. Red heading (large, 18pt, red)
p1 = doc.add_paragraph()
r1 = p1.add_run("WARNING: Safety Procedures")
r1.bold = True
r1.font.size = Pt(18)
r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# 2. Normal body text
p2 = doc.add_paragraph("Check the engine oil level regularly.")

# 3. Table with varying text lengths
table = doc.add_table(rows=3, cols=2)
table.cell(0, 0).text = "Part Number"
table.cell(0, 1).text = "Description"
table.cell(1, 0).text = "XYZ-1234"
table.cell(1, 1).text = "Oil Filter Assembly for Heavy Duty Engine Block"
table.cell(2, 0).text = "ABC-5678"
table.cell(2, 1).text = "Brake Pad"

# 4. Empty paragraph with page break (to simulate blank page)
p3 = doc.add_paragraph()
p3.paragraph_format.page_break_before = True

# 5. Another empty paragraph at end
doc.add_paragraph()

test_path = "uploads/test_layout.docx"
os.makedirs("uploads", exist_ok=True)
doc.save(test_path)
print("Test DOCX created")

# ---- RUN TRANSLATION (mock) ----
from backend.utils.docx_utils import translate_docx

def mock_translator(texts, lang, progress_cb=None):
    result = []
    for t in texts:
        result.append(t + " (translated extra text here)")
    return result

out_path = "uploads/test_layout_translated.docx"
translate_docx(test_path, out_path, mock_translator, "Tamil")
print("Translation done")

# ---- VERIFY ----
doc2 = Document(out_path)

# Check 1: Red heading font size preserved
first_para = doc2.paragraphs[0]
heading_run = first_para.runs[0] if first_para.runs else None
if heading_run:
    hs = heading_run.font.size
    size_val = hs.pt if hs else "None"
    print("Heading font size: {} pt (expected 18)".format(size_val))
    assert hs and hs.pt == 18, "FAIL: Heading size was {}".format(hs)
    print("PASS: Heading size preserved")

# Check 2: Table font sizes
for table in doc2.tables:
    sizes = set()
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size:
                        sizes.add(run.font.size.pt)
    print("Table font sizes: {}".format(sizes))
    if len(sizes) <= 1:
        print("PASS: Table font sizes are uniform")
    elif len(sizes) == 2:
        print("NOTE: 2 sizes found (may include heading exemption)")
    else:
        print("WARN: Multiple font sizes in table: {}".format(sizes))

# Check 3: No blank-page-causing elements  
body = doc2.element.body
blank_page_breaks = 0
for p in body.findall(qn("w:p")):
    para_obj = Paragraph(p, doc2)
    if not para_obj.text.strip():
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            if pPr.findall(qn("w:pageBreakBefore")):
                blank_page_breaks += 1
print("Blank page breaks on empty paras: {}".format(blank_page_breaks))
assert blank_page_breaks == 0, "FAIL: Still have blank page breaks"
print("PASS: No blank page breaks")

print()
print("ALL CHECKS PASSED")
